import re
import datetime
import openpyxl
from docx import Document
from pathlib import Path

def format_value(placeholder_text, raw_value):
    """Formats currency for [$...] placeholders, dates for datetime values."""
    if isinstance(raw_value, (datetime.datetime, datetime.date)):
        return raw_value.strftime("%m/%d/%Y")

    is_dollar = bool(re.match(r"^\[\$", placeholder_text.strip()))
    if not is_dollar:
        return str(raw_value)
    try:
        num = float(str(raw_value).replace(",", "").replace("$", ""))
    except ValueError:
        return str(raw_value)
    return f"${num:,.2f}"

def read_excel_mappings(excel_path):
    """Reads placeholders and values from an Excel file."""
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = None
    # Look for a sheet with 'field' in the name
    for name in wb.sheetnames:
        if "field" in name.lower():
            ws = wb[name]
            break
    
    if ws is None:
        raise ValueError(f"No sheet with 'field' in name found. Sheets: {wb.sheetnames}")

    mappings = []
    skipped = []
    for row in ws.iter_rows(min_row=1, values_only=True):
        if len(row) < 3:
            continue
        field_cell, value_cell = row[1], row[2]
        if not field_cell or not str(field_cell).strip():
            continue
        field_raw = str(field_cell).strip()
        if not re.search(r"\[", field_raw):
            continue
        if not value_cell or not str(value_cell).strip():
            skipped.append(field_raw)
            continue
        mappings.append((field_raw, format_value(field_raw, str(value_cell).strip())))

    # Sort by length descending to prevent partial replacements (e.g., [FIELD] vs [FIELD_1])
    mappings.sort(key=lambda x: len(x[0]), reverse=True)
    return mappings, skipped

def read_excel_records(excel_path, sheet_name=None, placeholder_row=1, header_row=2, data_start_row=3):
    """Reads a tabular export sheet where one row in the header block lists
    bracketed placeholders. Returns (records, placeholder_columns, header_row_values).

    records: list of dicts shaped {'mappings': [(placeholder, value), ...], 'row': raw_row_tuple}
    placeholder_columns: list of (col_index, placeholder_text)
    header_row_values: tuple of human-readable column headers from `header_row`
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    ws = None
    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        # Default to a sheet whose name starts with "Export" but isn't a pivot
        for name in wb.sheetnames:
            n = name.lower().strip()
            if n.startswith("export") and "pivot" not in n:
                ws = wb[name]
                break
        if ws is None:
            ws = wb[wb.sheetnames[0]]

    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < data_start_row:
        raise ValueError(f"Sheet '{ws.title}' has too few rows for bulk fill.")

    placeholder_row_vals = rows[placeholder_row - 1]
    header_row_vals = rows[header_row - 1]

    placeholder_columns = []
    for idx, cell in enumerate(placeholder_row_vals):
        if cell is None:
            continue
        text = str(cell).strip()
        if not text or "[" not in text or "]" not in text:
            continue
        placeholder_columns.append((idx, text))

    if not placeholder_columns:
        raise ValueError(
            f"No bracketed placeholders found in row {placeholder_row} of sheet '{ws.title}'."
        )

    records = []
    for raw_row in rows[data_start_row - 1:]:
        if not any(c is not None and str(c).strip() != "" for c in raw_row):
            continue

        mappings = []
        for col_idx, placeholder in placeholder_columns:
            if col_idx >= len(raw_row):
                continue
            val = raw_row[col_idx]
            if val is None:
                continue
            s = str(val).strip()
            if not s:
                continue
            mappings.append((placeholder, format_value(placeholder, val)))

        if not mappings:
            continue

        mappings.sort(key=lambda x: len(x[0]), reverse=True)
        records.append({"mappings": mappings, "row": raw_row})

    return records, placeholder_columns, header_row_vals

# NSA sheet fixed column indices (0-based). These columns have no bracket
# label in row 0 but carry data that maps to a template placeholder.
NSA_EXTRA_PLACEHOLDER_COLS = {
    "[date of service]": 28,   # 'DOS' column
    "[dispute ID]": 39,        # 'Payor_Claim' column — holds NSA dispute ID
}

# Template alias: templates use [tech] but header says [technologist]
NSA_TECH_ALIAS = "[tech]"
NSA_TECH_LABEL = "[technologist]"

# Row indices (0-based) for routing / filename metadata
NSA_PATIENT_COL   = 1    # 'Patient Name'
NSA_COMPS_COL     = 8    # 'Comps'
NSA_PROC_TYPE_COL = 7    # 'Category' (B&S / Pain)
NSA_DISPUTE_COL   = 39   # 'Payor_Claim' — NSA dispute ID


def read_excel_records_ps_tabular(excel_path, sheet_name=None):
    """NSA tabular PS reader.

    Sheet shape (Book1.xlsx 'Fields to Enter'):
      row 0  -> column headers; columns that hold template variables already
                carry the bracketed placeholder as their header label (e.g.
                '[provider]', '[carrier]', etc.).  The column index in row 0
                IS the data column for that placeholder.
      row 1+ -> one record per data row.

    Auto-detects bracketed placeholders from row 0.
    Adds '[date of service]' from the 'DOS' column (NSA_EXTRA_PLACEHOLDER_COLS).
    Adds '[tech]' alias for '[technologist]' so both template variants work.

    Returns list of dicts: {mappings, patient_name, dispute_id,
                            num_comps, procedure_type, row}.
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    ws = None
    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        for name in wb.sheetnames:
            if "field" in name.lower():
                ws = wb[name]
                break
        if ws is None:
            ws = wb[wb.sheetnames[0]]

    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise ValueError(f"Sheet '{ws.title}' has too few rows for NSA tabular fill.")

    header_row = rows[0]

    # Build placeholder -> col_idx from row 0 bracket labels
    auto_cols = {}
    for ci, cell in enumerate(header_row):
        if cell is None:
            continue
        text = str(cell).strip()
        if "[" in text and "]" in text:
            auto_cols[text] = ci

    # Merge extra fixed cols (DOS etc.)
    placeholder_cols = dict(auto_cols)
    for ph, ci in NSA_EXTRA_PLACEHOLDER_COLS.items():
        if ph not in placeholder_cols:
            placeholder_cols[ph] = ci

    # Add [tech] alias if [technologist] was detected
    if NSA_TECH_LABEL in placeholder_cols and NSA_TECH_ALIAS not in placeholder_cols:
        placeholder_cols[NSA_TECH_ALIAS] = placeholder_cols[NSA_TECH_LABEL]

    if not placeholder_cols:
        raise ValueError(
            f"No bracketed placeholders found in row 0 of sheet '{ws.title}'."
        )

    def _cell_value(raw_row, idx):
        if idx is None or idx >= len(raw_row):
            return None
        return raw_row[idx]

    def _cell_str(raw_row, idx):
        v = _cell_value(raw_row, idx)
        if v is None:
            return None
        s = str(v).strip()
        return s if s else None

    records = []
    for raw_row in rows[1:]:   # data starts at row index 1
        if not any(c is not None and str(c).strip() != "" for c in raw_row):
            continue
        if _cell_str(raw_row, NSA_PATIENT_COL) is None:
            continue

        mappings = []
        seen = set()
        for placeholder, col_idx in placeholder_cols.items():
            if placeholder in seen:
                continue
            val = _cell_value(raw_row, col_idx)
            if val is None:
                continue
            s = str(val).strip()
            if not s:
                continue
            mappings.append((placeholder, format_value(placeholder, val)))
            seen.add(placeholder)

        if not mappings:
            continue

        mappings.sort(key=lambda x: len(x[0]), reverse=True)

        num_comps = None
        raw_comps = _cell_value(raw_row, NSA_COMPS_COL)
        if raw_comps is not None:
            try:
                num_comps = int(raw_comps)
            except (ValueError, TypeError):
                try:
                    num_comps = int(float(raw_comps))
                except (ValueError, TypeError):
                    num_comps = None

        records.append({
            "mappings": mappings,
            "row": raw_row,
            "patient_name": _cell_str(raw_row, NSA_PATIENT_COL),
            "dispute_id": _cell_str(raw_row, NSA_DISPUTE_COL),
            "num_comps": num_comps,
            "procedure_type": _cell_str(raw_row, NSA_PROC_TYPE_COL),
        })

    return records


# UPM sheet fixed cols (0-based, no bracket label in row 0)
UPM_PATIENT_COL    = 0   # 'Patient Name'
UPM_CLAIM_TYPE_COL = 3   # 'ClaimType' → 'TDI' or 'NSA'
UPM_DISPUTE_COL    = 33  # 'Payor_Claim' — dispute/claim number for filename

# Alias map: template placeholder -> Excel header placeholder that holds same data.
# Used when template and Excel use different names for the same field.
UPM_PLACEHOLDER_ALIASES = {
    "[CMS – Public Use File Award Count]": "[CMS disputed claims]",  # col 52
    "[$CMS-PUF]": "[$CMS – PUF value]",                             # col 53
    "[cms to billed charges %]": "[cms to billed charges %]",       # identical, no-op
}


def read_excel_records_upm(excel_path, sheet_name=None):
    """UPM tabular reader.

    Sheet shape (Book12.xlsx 'Sheet1'):
      row 0  -> column headers; bracketed headers ARE the placeholder names
                at their data column index.
      row 1+ -> one record per row.

    Auto-detects bracketed placeholder columns from row 0.
    Routes each record by ClaimType column (col 3): 'TDI' or 'NSA'.

    Returns list of dicts: {mappings, patient_name, dispute_id, case_type, row}.
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    ws = None
    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        ws = wb[wb.sheetnames[0]]

    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        raise ValueError(f"Sheet '{ws.title}' has too few rows for UPM fill.")

    header_row = rows[0]

    # Auto-detect bracketed placeholder cols from row 0
    placeholder_cols = {}
    for ci, cell in enumerate(header_row):
        if cell is None:
            continue
        text = str(cell).strip()
        if "[" in text and "]" in text:
            placeholder_cols[text] = ci

    if not placeholder_cols:
        raise ValueError(
            f"No bracketed placeholders found in row 0 of sheet '{ws.title}'."
        )

    def _cell_value(raw_row, idx):
        if idx is None or idx >= len(raw_row):
            return None
        return raw_row[idx]

    def _cell_str(raw_row, idx):
        v = _cell_value(raw_row, idx)
        if v is None:
            return None
        s = str(v).strip()
        return s if s else None

    records = []
    for raw_row in rows[1:]:
        if not any(c is not None and str(c).strip() != "" for c in raw_row):
            continue
        if _cell_str(raw_row, UPM_PATIENT_COL) is None:
            continue

        mappings = []
        seen = set()
        for placeholder, col_idx in placeholder_cols.items():
            if placeholder in seen:
                continue
            val = _cell_value(raw_row, col_idx)
            if val is None:
                continue
            s = str(val).strip()
            if not s:
                continue
            mappings.append((placeholder, format_value(placeholder, val)))
            seen.add(placeholder)

        if not mappings:
            continue

        # Add alias entries so template placeholders with different names still fill
        existing_keys = {ph for ph, _ in mappings}
        alias_additions = []
        # Build reverse lookup: excel_placeholder -> value from mappings
        val_by_ph = {ph.lower(): v for ph, v in mappings}
        for tmpl_ph, excel_ph in UPM_PLACEHOLDER_ALIASES.items():
            if tmpl_ph in existing_keys:
                continue  # already present
            v = val_by_ph.get(excel_ph.lower())
            if v is not None:
                alias_additions.append((tmpl_ph, v))
        mappings = alias_additions + mappings
        mappings.sort(key=lambda x: len(x[0]), reverse=True)

        case_type = _cell_str(raw_row, UPM_CLAIM_TYPE_COL) or ''

        records.append({
            "mappings": mappings,
            "row": raw_row,
            "patient_name": _cell_str(raw_row, UPM_PATIENT_COL),
            "dispute_id": _cell_str(raw_row, UPM_DISPUTE_COL),
            "case_type": case_type.upper(),
        })

    return records


def read_excel_records_column_oriented(excel_path, sheet_name=None):
    """Reads a column-oriented sheet where column B holds field names
    (placeholders like [Procedure] or labels like 'Patient Name') and each
    subsequent column (C onward) is one record.

    Returns list of dicts: {mappings, patient_name, procedure, dispute_id,
                             num_comps, procedure_type}.
    """
    wb = openpyxl.load_workbook(excel_path, data_only=True)

    ws = None
    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        for name in wb.sheetnames:
            n = name.lower().strip()
            if "field" in n and ("replace" in n or "fill" in n or "enter" in n):
                ws = wb[name]
                break
        if ws is None:
            for name in wb.sheetnames:
                if "field" in name.lower():
                    ws = wb[name]
                    break
        if ws is None:
            raise ValueError(f"No field sheet found. Sheets: {wb.sheetnames}")

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError(f"Sheet '{ws.title}' is empty.")

    max_cols = max((len(r) for r in rows), default=0)
    if max_cols < 3:
        raise ValueError(
            f"Sheet '{ws.title}' has too few columns for column-oriented fill."
        )

    placeholder_rows = []
    patient_row = None
    procedure_row = None
    num_comps_row = None
    procedure_type_row = None

    for idx, row in enumerate(rows):
        if len(row) < 2 or row[1] is None:
            continue
        label = str(row[1]).strip()
        if not label:
            continue
        if "[" in label and "]" in label:
            placeholder_rows.append((idx, label))
            if re.sub(r"\s+", "", label).lower() == "[procedure]":
                procedure_row = idx
        label_lower = label.lower()
        if label_lower in ("patient name", "patient"):
            patient_row = idx
        elif "number of comps" in label_lower:
            num_comps_row = idx
        elif label_lower == "procedure type":
            procedure_type_row = idx

    if not placeholder_rows:
        raise ValueError(
            f"No bracketed placeholders found in column B of sheet '{ws.title}'."
        )

    records = []
    for col in range(2, max_cols):
        mappings = []
        for r_idx, placeholder in placeholder_rows:
            r = rows[r_idx]
            if col >= len(r):
                continue
            val = r[col]
            if val is None:
                continue
            s = str(val).strip()
            if not s:
                continue
            mappings.append((placeholder, format_value(placeholder, val)))

        if not mappings:
            continue

        mappings.sort(key=lambda x: len(x[0]), reverse=True)

        def _get_str(row_idx):
            if row_idx is None or col >= len(rows[row_idx]):
                return None
            v = rows[row_idx][col]
            return str(v).strip() if v is not None and str(v).strip() else None

        patient_name = _get_str(patient_row)
        procedure = _get_str(procedure_row)
        procedure_type = _get_str(procedure_type_row)

        num_comps = None
        if num_comps_row is not None and col < len(rows[num_comps_row]):
            raw = rows[num_comps_row][col]
            if raw is not None:
                try:
                    num_comps = int(raw)
                except (ValueError, TypeError):
                    num_comps = None

        dispute_id = None
        for ph, vv in mappings:
            if re.sub(r"\s+", "", ph).lower().strip("[]") == "disputeid":
                dispute_id = vv
                break

        records.append({
            "mappings": mappings,
            "patient_name": patient_name,
            "procedure": procedure,
            "dispute_id": dispute_id,
            "num_comps": num_comps,
            "procedure_type": procedure_type,
        })

    return records


def is_scs_procedure(procedure_text):
    """True if the procedure text references a Spinal Cord Stimulator."""
    if not procedure_text:
        return False
    s = str(procedure_text)
    if "spinal cord stimulator" in s.lower():
        return True
    if re.search(r"\bSCS\b", s):
        return True
    return False


def safe_filename_part(value, fallback="record"):
    """Sanitizes a value for use in a filename."""
    if value is None:
        return fallback
    if isinstance(value, (datetime.datetime, datetime.date)):
        s = value.strftime("%Y-%m-%d")
    else:
        s = str(value)
    s = re.sub(r"[^A-Za-z0-9._-]+", "_", s).strip("_")
    return s or fallback

def get_all_runs(paragraph):
    """Helper to extract all runs from a paragraph, including those in hyperlinks."""
    from docx.text.run import Run
    runs = []
    for child in paragraph._p:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            runs.append(Run(child, paragraph))
        elif tag == "hyperlink":
            for r_elem in child:
                r_tag = r_elem.tag.split("}")[-1] if "}" in r_elem.tag else r_elem.tag
                if r_tag == "r":
                    runs.append(Run(r_elem, paragraph))
    return runs

def replace_in_paragraph(paragraph, mappings):
    """Replaces placeholders across multiple runs in a paragraph."""
    runs = get_all_runs(paragraph)
    if not runs:
        return 0
    
    char_map = []
    for i, run in enumerate(runs):
        for ch in run.text:
            char_map.append((ch, i))
    
    if not char_map:
        return 0
        
    full_text = "".join(c for c, _ in char_map)
    full_text_low = full_text.lower()
    count = 0
    
    for placeholder, value in mappings:
        search_str = placeholder.lower()
        start_idx = 0
        while True:
            idx = full_text_low.find(search_str, start_idx)
            if idx == -1:
                break
            
            end_idx = idx + len(search_str)
            # Find which run this placeholder starts in
            run_index = char_map[idx][1]
            
            # Update char_map and full_text to reflect the replacement
            char_map = char_map[:idx] + [(ch, run_index) for ch in value] + char_map[end_idx:]
            full_text = full_text[:idx] + value + full_text[end_idx:]
            full_text_low = full_text_low[:idx] + value.lower() + full_text_low[end_idx:]
            
            count += 1
            start_idx = idx + len(value)
            
    if count == 0:
        return 0
        
    # Reconstruct run texts
    run_texts = {i: [] for i in range(len(runs))}
    for ch, i in char_map:
        run_texts[i].append(ch)
        
    for i, run in enumerate(runs):
        run.text = "".join(run_texts[i])
        
    return count

def fill_document(word_file, mappings):
    """Processes a Word document and replaces all placeholders."""
    doc = Document(word_file)
    total_replacements = 0
    
    # Process paragraphs
    for p in doc.paragraphs:
        total_replacements += replace_in_paragraph(p, mappings)
        
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total_replacements += replace_in_paragraph(p, mappings)
                    
    # Process headers and footers
    for section in doc.sections:
        headers_footers = [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer
        ]
        for hf in headers_footers:
            if hf:
                for p in hf.paragraphs:
                    total_replacements += replace_in_paragraph(p, mappings)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                total_replacements += replace_in_paragraph(p, mappings)

    return doc, total_replacements
